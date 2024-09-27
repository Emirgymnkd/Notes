import tkinter as tk  #GUI oluşturmak için kullanılan kütüphane.
from tkinter import filedialog, messagebox #filedialog dosya seçimi için pencere açar, messagebox ise bilgi mesajlarını göstermek için kullanılır.
from tkinter import ttk #görsel açıdan daha şık bileşenleri kullandıran ttk modülü. 
import os #işletim sistemiyle ilgili dosya ve klasör işlemlerini sağlar. dosya yolunu almak veya masaüstüne kaydetmek gibi.
import pandas as pd #veri analizlerinde kullanılan bir kütüphanedir. bu kodda excel dosyasını okuyup veri işlemek için kullanıyoruz.
import matplotlib.pyplot as plt #grafik çizmek için kullanılan bir kütüphanedir.
from openpyxl import load_workbook #excel dosyalarını okuma ve yazma işlemleri için tercih edilir.
from openpyxl.drawing.image import Image #excel dosyasına resim eklememizi sağlar.
from io import BytesIO #verileri bellekte saklamamızı sağlayan bir araçtır, bu kodda grafik resmini kaydediyoruz.
from docx import Document #word dosyasına ait bir kütüphanedir.
from docx.shared import Inches, Pt #word dosyasındaki nesneleri boyutlandırmak için kullanılır. Inches inç, Pt ise yazı boyutu cinsinden birimlerdir.
from docx.enum.text import WD_ALIGN_PARAGRAPH #worddeki paragraf hizalama ayarlarını yapmamızı sağlar.
from openpyxl.utils import get_column_letter #exceldeki sütun numaralarını harflerle eşleştirmek için kullanılır. Örneğin 1 numaralı sütun 'A' harfine denk gelir.
from openpyxl.utils import column_index_from_string #belirli desenlere göre metinleri incelemek ve bölmek için kullanılır.
import re

#bu fonksiyon, Y ekseni değerlerine göre grafikte kullanılacak nokta boyutlatını hesaplar.
#grafikteki noktaların büyük veya küçük görünmesi bu fonksiyon tarafından belirlenir.
def calculate_sizes(y_values, ratio=0.8, min_size=10, max_size=10):
    sizes = [min(max(min_size + (y * ratio), min_size), max_size) for y in y_values]
    return sizes
#eğer y_values = [1,2,3] ise ve oran(ratio) 0.8 ise, bu değerler her bir nokta boyutunu belirler. Boyutlar minimum ve maksimum 10 olacak şekilde sınırlandırılmıştır.


#bu fonksiyon bir balon grafiği oluşturur. balon grafikleri farklı boyutlara sahip noktaları göstermek için kullanılır.
#df, grafik için kullanılan verilerin DataFrame formatında olduğu anlamına gelir. aslında fonksiyonu çağırmış oluyoruz.
#x_column ve y_column, x ve y eksenindeki değerleri, size_values nokta boyutlarını (yukarıdaki calculate_sizes fonksiyonu ile hesaplanan değer) ifade eder.
#title: grafiğin başlığını açıklar.

def create_bubble_chart(df, x_column, y_column, size_values, title):
    #df[x_column] ve df[y_column], grafikte X ve Y eksenine karşılık gelen sütunları temsil eder. 
    # s=size_values, her bir noktanın boyutunu belirler.
    # calculate_sizes() fonksiyonundan gelen nokta büyüklükleridir. 
    #alpha, noktanın opaklığını belirler. 0-1 arası değerler alır.
    #color renkler, marker noktanın şeklini, edgecolors nokta kenarlık rengi, linewidths kenar kalınlığını ayarlar.
    plt.figure(figsize=(3.80, 3))  #bu komut yeni bir grafik oluşturur. boyutları inç cinsinden giriyoruz. (genişlik,yükseklik)
    plt.scatter(df[x_column], df[y_column], s=size_values, alpha=0.6, color='blue', marker='o', edgecolors='black' , linewidths=1.1)  
    plt.title(title, fontsize=12) #grafiğin başlığını ayarlar. title, bu fonksiyona parametre olarak geçirilir ve 12 punto büyüklüğünde yazılır.
    plt.xlabel(f"{x_column} Koordinatı", fontsize=10)
    plt.ylabel(f"{y_column} Koordinatı", fontsize=10) #eksen etiketlerini ayarlar. 10 puntodur.
    plt.tight_layout() #grafik içerisindeki öğelerin sıkışık görünmesini engeller. otomatik düzenleme yapar.
    imgdata = BytesIO() #grafik oluşturulduktan sonra bu  komut sayesinde grafik PNG formatında kaydedilir ve saklanır.
    plt.savefig(imgdata, format='png', dpi=300) #bellekte imgdata ile saklanır, dpi=300 yüksek çözünürlüklü bir grafik sağlar (dots per inch), görüntünün piksel yoğunluğunu temsil eder.
    imgdata.seek(0) #bellekte tutulan veriyi başa sarar. bu, grafiği başka bir yerden okumadan önce bellekteki konumu sıfırlamak içindir.
    plt.close() #grafiği kapatır.
    return imgdata #grafiği imgdata olarak geri döndürür. Bu veri daha sonra excel ve word dosyalarına resim olarak eklenir.

#bu fonksiyon, excel hücre konumlarını böler. Örneğin 'A1' gibi bir hücre konumunu harf kısmını ('A') ve sayı kısmını ('1') olarak döndürür.

def split_cell_position(cell_position):
    match = re.match(r"([A-Z]+)([0-9]+)", cell_position, re.I)
    if match:
        items = match.groups()
        return items[0], int(items[1]) 
    return None, None

#bu fonksiyon, her bir excel sayfası için balon grafiklerini oluşturur ve bu grafikleri excel ve word dosyasına ekler.
#writer excel dosyasına yazmak için kullanılırken, document word dosyasına yazmak için kullanılır.

def add_charts_to_excel_and_word(writer, df, sheet_name, document):
    workbook = writer.book
    chart_sheet = workbook.create_sheet(title=f"{sheet_name}_Charts")
    
    
    positions = {
    'FX': ('A1', 'R40'),  
    'FY': ('U1', 'AL40'),
    'FZ': ('AM1', 'BL40')
}
    
    # Her sayfanın FX, FY, FZ sütununa göre grafik oluşturur.
    for force, (start_cell, end_cell) in positions.items():
        df['FX'] = pd.to_numeric(df['FX'], errors='coerce').fillna(0)
        df['FY'] = pd.to_numeric(df['FY'], errors='coerce').fillna(0)
        df['FZ'] = pd.to_numeric(df['FZ'], errors='coerce').fillna(0)

        if force == 'FX':
            size_values = calculate_sizes(abs(df['FX']), ratio=0.5, min_size=5, max_size=500)  # Bu sayfa için FX sütunu
        elif force == 'FY':
            size_values = calculate_sizes(abs(df['FY']), ratio=0.5, min_size=5, max_size=500)  # Bu sayfa için FY sütunu
        elif force == 'FZ':
            size_values = calculate_sizes(abs(df['FZ'])/100, ratio=0.5, min_size=5, max_size=500)  # Bu sayfa için FZ sütunu


        # Excel'e grafiği ekle
        imgdata = create_bubble_chart(df, 'X', 'Y', size_values, f'{sheet_name} - {force} Kuvveti') #balon grafiği oluşturur. df=grafik için kullanılacak veriler , X,Y = x ve y sütununda gösterilecek sütunlar, size_values = balonların boyutlarını belirler, f'{sheet_name} - {force} Kuvveti' , grafik başlığını belirler.
        img = Image(imgdata) #grafik PNG formatında kaydedilir ve bir resime dönüştürülür.
        chart_sheet.add_image(img, start_cell)#grafiği excel sayfasına ekler, parantez içi ise grafiği seçer ve yerleştirilmeye başlayacağı hücre konumunu belirtir.

        # Word dosyasına grafiği ekle
        document.add_paragraph(f'{sheet_name} - {force} Kuvveti Dağılımı') #worde dosya ekler ve paragrafın içeriğini başlığını parantez içerisinde belirtmen istenir.
        document.add_picture(BytesIO(imgdata.getvalue()), width=Inches(6)) #bellekte tutulan grafiği ekler, getvalue() sayesinde o veriyi almış oluyoruz. width ile de grafiğin genişliğini belirliyoruz.
        document.add_page_break() #sayfa sonuna gelir ve yeni bir sayfa başlatır.

#bu fonksiyon seçilen excel dosyasını işler. Belirtilen sayfaları okur, temizler ve her bir sayfa için grafikleri tanımlandığı gibi oluşturur.
#aynı zamanda excel ve word dosyalarını oluşturur.

def process_file(file_path): #seçilen dosyayı bu fonksiyon işler, grafikleri ekler. bu satırlar exceldeki 'joint reactions' ve 'point object connectivity' sayfalarını okur ve bu sayfalardaki veirleri dataframe formatında yükler.
    joint_reactions_df = pd.read_excel(file_path, sheet_name='Joint Reactions', skiprows=0) #skiprows ifadesi ilk satırdan itibaren tüm satırların okunmasını sağlar.
    point_object_connectivity_df = pd.read_excel(file_path, sheet_name='Point Object Connectivity', skiprows=0) #aynı şekilde 'point object connectivity' sayfasını okur.

    
    joint_reactions_df_cleaned = joint_reactions_df.dropna(how='all')  #tamamen boş olan satıları kaldırır.
    joint_reactions_df_cleaned.columns = joint_reactions_df_cleaned.iloc[0] #ilk satırı sütun başlıkları olarak ayarlar.
    joint_reactions_df_cleaned = joint_reactions_df_cleaned.drop(0)  #ilk satırı kaldırır. buradaki drop komutu belirli bir satırı veya sütunu silmek için kullanılan bir komuttur.
    

    """sütun silmek isteseydim:
    df.drop('Sütun_Adı', axis=1, inplace=True) kalıbında bir komut yazmalıydım. axis=1 diyerek sütun üzerinde işlem gerçekleştireceğimi belirtmiş oluyorum. axis= 0 yazsaydım satır olarak algılayacaktı. örnek:
    df.drop('Age', axis=1, inplace=True)
    burada inplace de dataframe' in orijinalinde bir değişiklik yapmasını (true), yapmamasını (false) belirtiyorum"""



    # Clean the 'Point Object Connectivity' DataFrame similarly
    point_object_connectivity_df_cleaned = point_object_connectivity_df.dropna(how='all') #yine boş satırları kaldırıyorum.
    point_object_connectivity_df_cleaned.columns = point_object_connectivity_df_cleaned.iloc[0]  #ilk satırları sütun başlıkları olarak ayarlıyorum.
    point_object_connectivity_df_cleaned = point_object_connectivity_df_cleaned.drop(0)  # ilk satırı kaldırır.

    # Extract unique 'Output Case' values from 'Joint Reactions'
    output_cases = joint_reactions_df_cleaned['Output Case'].unique() #'output case' sütunundaki benzersiz değerleri belirler.

    # Filtering out NaN from the 'Output Case' unique values
    output_cases = [case for case in output_cases if pd.notna(case)] # output_cases verisi içerisindeki elemanları (case) kontrol eder ve sadece NaN olmayan elemanları yeni bir listeye dahil eder. Boş hücrelerden kurtuluyorum kısacası.

    # Step a: Filter 'Joint Reactions' data by 'Output Case'
    result_dfs = {} #işlenen verilerin saklanacağı bir sözlük oluşturur.
    for case in output_cases:
        # a. Filter the 'Joint Reactions' data by the current 'Output Case'
        filtered_joint_reactions = joint_reactions_df_cleaned[joint_reactions_df_cleaned['Output Case'] == case] #joint_reactions_df_cleaned verileri içerisinde 'output case' sütunundaki verilerin 'case' olup olmadığını kontrol edip filtreliyor.

        # b. Extract unique 'Label' values from this filtered data
        unique_labels = filtered_joint_reactions['Label'].unique() #label sütunundaki benzersiz değerleri tespit eder.

        # c. For each unique 'Label', find the corresponding coordinates in 'Point Object Connectivity'
        coords = point_object_connectivity_df_cleaned[point_object_connectivity_df_cleaned['PointBay'].isin(unique_labels)] #isin komutu sayesinde 'point_object_connectivity_df_cleaned' verileri içinde 'point bay' sütunundaki veriler ile 'unique_labels' verilerini kıyaslıyoruz.

        # d. Merge the filtered 'Joint Reactions' data with the coordinates data
        merged_data = filtered_joint_reactions.merge(coords, left_on='Label', right_on='PointBay', how='left') #bu veriler 'label' ve 'point bay' sütununa göre yerleştirilir.

        # Split merged data by 'Step Type'
        step_type_groups = merged_data.groupby('Step Type') #'step type' sütununa göre birleştirme(gruplandırma) yapılır.

        # Save each group as a DataFrame
        for step_type, group in step_type_groups:
            result_dfs[f"{case}_Step_{step_type}"] = group #gruplanmış verileri yukarıda oluşturduğumuz sözlük içerisinde saklıyoruz. bunu yazmadığımız taktirde verileri kullanamaz, sonuçları kaydedemem. bu yüzden burası kritik.
            

    # Filter DESG0, DESG1, DESG3 and store them in separate dataframes
    desg_cases = {} #DESG tablolar için ayrı bir sözlük oluşturur.
    for desg in ['DESG0', 'DESG1', 'DESG3']:
        desg_cases[desg] = joint_reactions_df_cleaned[joint_reactions_df_cleaned['Output Case'] == desg]
        
        # Merge DESG tables with Point Object Connectivity coordinates
        merged_desg = desg_cases[desg].merge(point_object_connectivity_df_cleaned, left_on='Label', right_on='PointBay', how='left') #desg tabloları da koordinatlar ile birleştirilir. (merge())
        desg_cases[desg] = merged_desg  # Update the DESG dataframes with merged coordinates

    # Save the result to the desktop
    desktop_path = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop') #masaüstü yolu belirlenir. (os.path.expanduser('~')), kullanıcının ev dizinini yani masaüstüne ulaşacak yolu bulmak için kullanılır. ~ işareti ev dizini anlamına gelmekte.
    file_name = os.path.splitext(os.path.basename(file_path))[0] + " - filtrelenmiş.xlsx" #path = dosya yollarını yönetmek için kullanılır. splitext ise dosya adını ve uzantısı arasına boşluk ekler. bu boşluğa dosya ismi geleceği için gereklidir.
    #basename ise sadece dosya adını alır. tam dosya yolunu vermeme rağmen sadece dosyanın ismini seçer.
    #[0] sayesinde ise splitext ile dosya adı ve dosya uzantısı olarak böldüğümüz kısmın ilkini, yani dosya adını seçmemizi sağlar, uzantıyı almaz. örnek:
    dosya_adı_ve_uzantı = os.path.splitext('document.xlsx')
    # + filtrelenmiş kısmını da seçtiğimiz dosya ismine eklemesini yapıyoruz.
    output_file_path = os.path.join(desktop_path, file_name) #dosya yolunu belirlemiş oluyoruz.

    # Create a Word document
    word_file_name = os.path.splitext(os.path.basename(file_path))[0] + " - Grafikler.docx" #yukarıdakiyle aynı sistem word için gerçekleşiyor
    word_file_path = os.path.join(desktop_path, word_file_name)
    document = Document() #yeni bir word belgesi başlatılıyor.

    # Writing dataframes and adding charts to Excel
    with pd.ExcelWriter(output_file_path, engine='openpyxl') as writer: #bu döngü verilerin excele yazılmasını başlatır.
        # Write result dataframes
        for name, df in result_dfs.items(): #döngüyle beraber veriler excele aktarılır.
            df.to_excel(writer, sheet_name=name, index=False)

        # Write DESG cases with coordinates merged
        for desg_name, desg_df in desg_cases.items(): #aynı işlem desg ler için yapılır.
            desg_df.to_excel(writer, sheet_name=f"{desg_name}_filtered", index=False)

        # Add bubble charts to Excel and Word for each sheet's FX, FY, FZ values
        for name, df in result_dfs.items():
            add_charts_to_excel_and_word(writer, df, name, document) #excel ve word dosyasına grafik eklenir.

        # Add bubble charts for DESG cases
        for desg_name, desg_df in desg_cases.items(): #desg tabloları da aynı şekilde eklenir.
            add_charts_to_excel_and_word(writer, desg_df, desg_name, document)
            

    # Save Word document
    document.save(word_file_path) #kaydet.

    # Return the saved file path for notification
    return output_file_path, word_file_path #dosya yolunu döndürür.

#bu fonksiyon, işlemin başarıyla tamamlandığını gösteren bir pop-up penceresi açar
# #kullanıcıya excel ve word dosyalarının nerede kaydedildiğini gösterir.


def show_popup(output_file_paths):
    excel_file_path, word_file_path = output_file_paths #output_file_paths, iki farklı dosya yolunu içeren bir listedir.
    popup = tk.Toplevel() #bu komut, ana pencerenin üstüne başka bir pencere açmak için kullanılır.
    popup.title("İşlem Tamamlandı")
    
    # pop-up'ın ölçülerini belirler.
    popup_width, popup_height = 450, 250            #genişlik 250 ve yükseklik 450 olarak ayarlanıyor.
    popup.geometry(f"{popup_width}x{popup_height}") #bu verileri alarak oluşturuyor.

    # Get screen dimensions for centering the popup
    screen_width = popup.winfo_screenwidth() #kullanılan ekranın genişliğini alır.
    screen_height = popup.winfo_screenheight() #kullanılan ekranın yüksekliğini alır

    # Calculate the center coordinates for the popup
    center_y = int(screen_height / 2 - popup_height / 2) #pencerenin ortalanması için y konumunu hesaplar.

    # Center the popup
    popup.geometry(f"{popup_width}x{popup_height}+{center_x}+{center_y}") #pencerenin tam ekranın ortasına yerleştirilmesini sağlar.

    # Create label in the popup
    message = (f"İşlem başarıyla tamamlandı.\n\n"
               f"Excel Dosyası: '{excel_file_path}'\n"
               f"Word Dosyası: '{word_file_path}'\n"
               "konumlarına kaydedildi.") #ekrana gelecek olan yazılar.
    label = tk.Label(popup, text=message, font=("Helvetica", 12, "bold"), fg="black", justify="center") #bu mesajı gösterecek bir etiket(label) oluşturur. bu labelin bilgileri parantez içerisine girilir.
    label.pack(pady=40) #yukarıdan 40 piksel boşluk bırakılır.

    # Create an "OK" button to close the popup
    ok_button = tk.Button(popup, text="Tamam", command=popup.destroy, font=("Helvetica", 12, "bold")) #tk.button bir buton oluştururken, butona tıklandığında popup.destroy komutuyla pencere kapatılır.
    ok_button.pack(pady=10)

#kullanıcının dosyayı seçmesine olanak tanır ve seçilen dosya yolunu bir etikette gösterir.
def select_file(): #dosya seçme fonksiyonunu çalıştırıyoruz.
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")]) #filedialog = tkinter kütüphanesinin bir parçası olup dosya seçiminde kullanılır.
    #askopenfilename ise açılan pencereden kullanıcıdan bir dosya seçmesini ister, devamında ise excel dosyası ve xlsx formatında olması gerektiğini parantez içinde belirtir.
    if file_path:
        apply_button.config(state=tk.NORMAL) #aşağıda 'uygula' butonunu tanımladığımız apply_button'ı config komutuyla konfigüre ediyor, durumunu tk kütüphanesiyle DİSABLED'dan NORMAL'e yani aktif hale getiriyoruz. (eğer kullanıcı dosya seçtiyse tabi.)
        selected_file_label.config(text=f"Seçilen Dosya: {file_path}", fg="black") #'selected_file_label' olarak tanımladığımız 'seçilen dosya' yazısını kişiselleştiriyoruz ve ekranda görünür yapıyoruz.
        global selected_file #globalleştirmek, bu değişkeni hem fonksiyon içinde hem fonksiyon dışında kullanılabilir yapılmasını sağlar. local(yerel)'den (yani sadece fonksiyon içinde çalışma durumundan) kurtarmış oluyoruz.
        selected_file = file_path #seçilen dosyanın bu dosya yolunda olduğunu söylüyoruz.
    else:
        selected_file_label.config(text="Dosya seçimi yapmadınız.", fg="red") #eğer seçilmediyse bunu döndürmesini söylüyoruz. burada "fg"=foreground (ön plan) anlamına gelmekte olup, tkinterda kullanılan bir widgettır.

#seçilen dosyayı işler ve grafiklerin oluşturulmasını sağlar. işlem tamamlandığında show_popup() fonksiyonunu çağırarak bir bildirim gösterir.
def apply_changes():
    if selected_file:
        output_paths = process_file(selected_file)
        show_popup(output_paths)  # Show popup when processing is complete


root = tk.Tk() #tkinter ana penceresini oluşturur.
root.title("Dosya Seçimi ve İşleme") #pencerenin başlığını belirler.

# Set the window size (3 times larger)
window_width, window_height = 600, 400
root.geometry(f"{window_width}x{window_height}")

# Get screen dimensions
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Calculate the center coordinates for the window
center_x = int(screen_width/2 - window_width/2)
center_y = int(screen_height/2 - window_height/2)


root.geometry(f"{window_width}x{window_height}+{center_x}+{center_y}") #pencerenin boyutu ve konumu ayarlanır.


root.configure(bg='white')#pencerenin arka plan rengini beyaz yapar.

# kullanılacak yazının tipi, boyutu ve kalınlığı ayarlanır.
font_settings = ("Helvetica", 14, "bold")

# File selection button
select_button = tk.Button(root, text="Dosya Seç", command=select_file, font=font_settings)
select_button.pack(pady=20)

# Apply button (initially disabled)
apply_button = tk.Button(root, text="Uygula", command=apply_changes, font=font_settings, state=tk.DISABLED)
apply_button.pack(pady=20)

# Label to display selected file
selected_file_label = tk.Label(root, text="Seçilen Dosya:", font=font_settings, bg='white')
selected_file_label.pack(pady=10)

# Label to show process result
result_label = tk.Label(root, text="", font=font_settings, bg='white', fg='black')
result_label.pack(pady=10)


root.mainloop() #bu komut, GUI'yisürekli çalışır halde tutar. Kullanıcı bir butona basana kadar pencere açık kalır.